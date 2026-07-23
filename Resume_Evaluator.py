import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field
import streamlit as st

# Pehle Streamlit Secrets check karo
api_key = st.secrets.get("GROQ_API_KEY")

# Agar local machine hai to env variable use karo
if api_key is None:
    api_key = os.getenv("GROQ_API_KEY")

if not api_key:
    raise ValueError("Groq API key not found.")

client=Groq(api_key=my_api_key)
model="llama-3.3-70b-versatile"

# Default Job Description (Agar koi JD pass na ho to ye use hogi)
job_description = None

# -------------------- Job Description Parser -------------------- #

def parse_job_description(job_description):

    # AI ko role aur task batane ke liye system prompt
    system_prompt = f"""
    You are an expert HR assistant.
    Your job is to analyze job descriptions and extract structured information from them.

    Return ONLY valid JSON matching this schema:
    {jobd_schema}

    IMPORTANT:

    Do NOT return the schema itself.

    Do NOT return fields like "properties", "title" or "type".

    Fill the schema with actual information extracted from the job description.

    If minimum experience is not mentioned, return null.

    If information for a list is missing, return an empty list.

    Do not invent information.
    """

    # User prompt bana rahe hain
    user_prompt = f"Analyze the following job description:\n{job_description}"

    # System message
    message_system = {
        "role": "system",
        "content": system_prompt
    }

    # User message
    message_user = {
        "role": "user",
        "content": user_prompt
    }

    # Messages list
    messages = [message_system, message_user]

    # JSON output force kar rahe hain
    response_format = {
        "type": "json_object"
    }

    # LLM call
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        response_format=response_format
    )

    # JSON response ko Python dictionary bana rahe hain
    job_data = json.loads(
        response.choices[0].message.content
    )

    # Validate karke object return
    return JobD(**job_data)

# Job Description ke liye ek structured data model bana rahe hain
class JobD(BaseModel):
    role: str  # Job role store karega (e.g., Data Scientist, Software Engineer)
    required_skills: list[str]  # Job ke liye mandatory skills ki list
    preferred_skills: list[str]  # Achha hoga agar candidate ke paas ye skills bhi ho
    minimum_experience: float | None  # Minimum experience (years me), agar mention na ho to None rahega
    education_requirements: list[str]  # Required education qualifications ki list
    responsibilities: list[str]    # Job ki responsibilities/tasks ki list

# Model ka JSON Schema generate kar rahe hain
# Ye AI ko batata hai ki output kis structure/format me hona chahiye
jobd_schema = JobD.model_json_schema()


# AI ko role aur task batane ke liye system prompt define kar rahe hain
system_prompt = f"""
You are an expert HR assistant.
Your job is to analyze job descriptions and extractstructured information from them.
Return ONLY valid JSON matching this schema:{jobd_schema}
IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.
If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""

user_prompt = f"""Analyze the following job description:{job_description}"""

message_system = {                      # System message dictionary bana rahe hain
    "role": "system",                  # AI ko system role assign kar rahe hain
    "content": system_prompt           # System instructions pass kar rahe hain
}

message_user = {                       # User message dictionary bana rahe hain
    "role": "user",                    # AI ko user role assign kar rahe hain
    "content": user_prompt             # User ka actual input/JD pass kar rahe hain
}

response_format = {                    # Response ka format define kar rahe hain
    "type": "json_object"              # AI se sirf JSON output mang rahe hain
}

messages = [message_system, message_user]  # Dono messages ko ek list me combine kar rahe hain

job = None



# Resume aur Job matching result ka structure define kar rahe hain
class MatchResult(BaseModel):
    score: float                  # Resume ka overall matching score
    details: dict                 # Matching ki detailed information

# Candidate ke work experience ka structure define kar rahe hain
class Experience(BaseModel):
    company: str | None = None    # Company ka naam
    role: str | None = None       # Job role
    duration: str | None = None   # Job duration
    description: str | None = None # Work ka short description
    skills_used: list[str] = []   # Is experience me use hui skills

# Resume ka complete structure define kar rahe hain
class Resume(BaseModel):
    name: str | None = None       # Candidate ka naam
    email: str | None = None      # Candidate ki email ID
    phone: str | None = None      # Candidate ka phone number

    total_experience_years: float | None = None  # Total experience (years me)

    skills: list[str] = []        # Candidate ki skills
    experiences: list[Experience] = []  # Sare work experiences
    education: list[str] = []     # Educational qualifications
    projects: list[str] = []      # Candidate ke projects
    certifications: list[str] = [] # Certifications ki list

resume_schema = Resume.model_json_schema()  # Resume model ka JSON schema generate kar rahe hain



def final_score(job, resume):
    match_schema = MatchResult.model_json_schema()  # MatchResult model ka JSON schema generate kar rahe hain

    prompt = f"""
    You are an HR recruiter.
    Compare the candidate's resume with the job description.
    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}  # Job object ko JSON format me convert kar rahe hain

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}  # Resume object ko JSON format me convert kar rahe hain
    Return JSON matching this schema:
    {match_schema}
    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """

    message = {
        "role": "user",          # User message create kar rahe hain
        "content": prompt        # AI ko prompt bhej rahe hain
    }

    messages = [message]         # Message ko list me store kar rahe hain

    response_format = {
        "type": "json_object"    # AI se sirf JSON response mang rahe hain
    }

    response = client.chat.completions.create(
        model=model,                     # Selected model use kar rahe hain
        messages=messages,               # Prompt send kar rahe hain
        response_format=response_format  # JSON output force kar rahe hain
    )

    data = json.loads(response.choices[0].message.content)  # JSON response ko Python dictionary me convert kar rahe hain

    return MatchResult(**data)  # Dictionary ko MatchResult object me convert karke return kar rahe hain




def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.
    Extract information from the resume based on its meaning,
    not only based on exact section headings.
    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.
    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:
    {resume_schema} 

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """

    user_prompt = f"""Parse the following resume:{resume_text}"""

    message_system = {
        "role": "system",          # System instructions set kar rahe hain
        "content": system_prompt
    }

    message_user = {
        "role": "user",            # User ka resume input bhej rahe hain
        "content": user_prompt
    }

    messages = [message_system, message_user]  # Dono messages ko list me combine kar rahe hain

    response_format = {
        "type": "json_object"      # AI se sirf JSON response mang rahe hain
    }

    response = client.chat.completions.create(
        model=model,                      # Selected model use kar rahe hain
        messages=messages,                # System + User messages bhej rahe hain
        response_format=response_format   # JSON output force kar rahe hain
    )

    raw_output = response.choices[0].message.content  # AI ka raw JSON response le rahe hain

    data = json.loads(raw_output)# JSON string ko Python dictionary me convert kar rahe hain

    # Resume level lists
    data["skills"] = data.get("skills") or []
    data["education"] = data.get("education") or []
    data["projects"] = data.get("projects") or []
    data["certifications"] = data.get("certifications") or []
    data["experiences"] = data.get("experiences") or []

    # Experience level lists
    for exp in data["experiences"]:
        exp["skills_used"] = exp.get("skills_used") or []  

    resume = Resume(**data)  # Dictionary ko Resume object me validate karke convert kar rahe hain

    return resume  # Parsed Resume object return kar rahe hain



from pypdf import PdfReader
from docx import Document

def read_pdf(file_path):
    reader = PdfReader(file_path)          # PDF file ko open kar rahe hain
    text = ""                              # Saara extracted text store hoga

    for page in reader.pages:              # PDF ke har page par loop chala rahe hain
        page_text = page.extract_text()    # Current page ka text extract kar rahe hain

        if page_text:                      # Agar page me text mila
            text += page_text + "\n"       # To usse final text me add kar do

    return text                            # Complete PDF text return kar rahe hain


def read_docx(file_path):
    document = Document(file_path)         # DOCX file ko open kar rahe hain
    text = ""                              # Extracted text store hoga

    for paragraph in document.paragraphs:  # Sabhi paragraphs par loop chala rahe hain
        if paragraph.text.strip():         # Empty paragraph skip kar rahe hain
            text += paragraph.text + "\n"  # Paragraph text add kar rahe hain

    for table in document.tables:          # Resume me agar tables hain
        for row in table.rows:             # Har row par loop
            for cell in row.cells:         # Har cell par loop
                if cell.text.strip():      # Empty cell skip kar rahe hain
                    text += cell.text + "\n"  # Cell ka text add kar rahe hain

    return text                            # Complete DOCX text return kar rahe hain


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":     # Agar PDF hai
        return read_pdf(file_path)             # PDF reader call karo

    elif file_path.suffix.lower() == ".docx":  # Agar DOCX hai
        return read_docx(file_path)            # DOCX reader call karo

    else:
        return None                            # Unsupported file type


# Main execution starts here
if __name__ == "__main__":

    # Resume folder ka path le rahe hain
    resume_folder = Path("resumes")

    all_results = []  # Sab candidates ke results store honge

    for file_path in resume_folder.iterdir():  # Folder ki har file par loop

        # Sirf PDF aur DOCX files process karni hain
        if file_path.suffix.lower() not in [".pdf", ".docx"]:
            continue

        print("\nProcessing:", file_path.name)  # Current file ka naam print kar rahe hain

        resume_text = read_resume(file_path)  # Resume ka text extract kar rahe hain

        parsed_resume = parse_resume(resume_text)  # LLM Call 1 → Resume parse kar rahe hain

        time.sleep(5)  # API rate limit avoid karne ke liye wait

        result = final_score(job, parsed_resume)  # LLM Call 2 → Resume aur JD compare kar rahe hain

        time.sleep(5)  # Next API request se pehle wait

        print("Score:", result.score)  # Candidate ka score print kar rahe hain

        all_results.append({
            "name": parsed_resume.name,      # Candidate ka naam
            "score": result.score,           # Matching score
            "details": result.details        # Matching details
        })

    # Candidates ko highest score ke hisaab se sort kar rahe hain
    all_results.sort(
        key=lambda candidate: candidate["score"],
        reverse=True
    )

    top_2 = all_results[:2]      # Top 2 candidates
    worst_2 = all_results[-2:]   # Lowest 2 candidates

    # Output Window
    print("\n" + "=" * 80)
    print("🏆 TOP 2 CANDIDATES")
    print("=" * 80)

    for i, candidate in enumerate(top_2, start=1):
        details = candidate["details"]

        print(f"\n🥇 Rank #{i}")
        print("-" * 80)
        print(f"👤 Name                 : {candidate['name']}")
        print(f"📊 Match Score          : {candidate['score']}%")
        print(f"✅ Matching Skills      : {', '.join(details.get('matching_skills', []))}")
        print(f"❌ Missing Skills       : {', '.join(details.get('missing_important_skills', []))}")
        print(f"💼 Experience Match     : {details.get('experience_requirement_met')}")
        print(f"📝 Final Verdict        : {details.get('final_verdict')}")
        print("-" * 80)

    print("\n\n" + "=" * 80)
    print("❌ LOWEST 2 CANDIDATES")
    print("=" * 80)

    for i, candidate in enumerate(worst_2, start=1):
        details = candidate["details"]

        print(f"\n🔻 Rank #{i}")
        print("-" * 80)
        print(f"👤 Name                 : {candidate['name']}")
        print(f"📊 Match Score          : {candidate['score']}%")
        print(f"✅ Matching Skills      : {', '.join(details.get('matching_skills', []))}")
        print(f"❌ Missing Skills       : {', '.join(details.get('missing_important_skills', []))}")
        print(f"💼 Experience Match     : {details.get('experience_requirement_met')}")
        print(f"📝 Final Verdict        : {details.get('final_verdict')}")
        print("-" * 80)